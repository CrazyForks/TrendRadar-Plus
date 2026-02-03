# coding=utf-8
"""
Document Parser Module

Provides parsers for various document formats:
- Markdown (.md)
- PDF (.pdf)
- Word (.docx)
"""

import io
import re
from typing import Tuple, Optional
from pathlib import Path


class DocumentParseError(Exception):
    """Document parsing error."""
    pass


def parse_markdown(content: bytes, filename: str = "") -> Tuple[str, str, str]:
    """
    Parse Markdown content to HTML.
    
    Args:
        content: Raw file content
        filename: Original filename (used for title extraction)
        
    Returns:
        Tuple of (title, digest, html_content)
    """
    try:
        import markdown
        from markdown.extensions import fenced_code, tables, toc
    except ImportError:
        raise DocumentParseError("缺少 markdown 库，请安装: pip install markdown")
    
    # Decode content
    try:
        text = content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            text = content.decode('gbk')
        except UnicodeDecodeError:
            raise DocumentParseError("无法解码文件，请确保文件编码为 UTF-8")
    
    # Extract title from first H1 or filename
    title = ""
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            title = line[2:].strip()
            break
    
    if not title and filename:
        # Use filename without extension as title
        title = Path(filename).stem
    
    # Extract digest from first paragraph
    digest = ""
    in_code_block = False
    for line in lines:
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        line = line.strip()
        if line and not line.startswith('#') and not line.startswith('---'):
            # Remove markdown formatting for digest
            digest = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # bold
            digest = re.sub(r'\*([^*]+)\*', r'\1', digest)    # italic
            digest = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', digest)  # links
            digest = re.sub(r'`([^`]+)`', r'\1', digest)      # inline code
            break
    
    # Convert to HTML
    md = markdown.Markdown(extensions=[
        'fenced_code',
        'tables',
        'toc',
        'nl2br',
    ])
    html_content = md.convert(text)
    
    return title[:64], digest[:200], html_content


def parse_pdf(content: bytes, filename: str = "") -> Tuple[str, str, str]:
    """
    Parse PDF content to HTML.
    
    Args:
        content: Raw file content
        filename: Original filename
        
    Returns:
        Tuple of (title, digest, html_content)
    """
    try:
        import pdfplumber
    except ImportError:
        raise DocumentParseError("缺少 pdfplumber 库，请安装: pip install pdfplumber")
    
    try:
        # Open PDF from bytes
        pdf_file = io.BytesIO(content)
        
        with pdfplumber.open(pdf_file) as pdf:
            if len(pdf.pages) == 0:
                raise DocumentParseError("PDF 文件为空")
            
            # Extract text from all pages
            all_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text)
            
            if not all_text:
                raise DocumentParseError("无法从 PDF 中提取文本")
            
            full_text = '\n\n'.join(all_text)
            
            # Extract title from first line or filename
            lines = full_text.split('\n')
            title = ""
            for line in lines:
                line = line.strip()
                if line and len(line) < 100:  # Reasonable title length
                    title = line
                    break
            
            if not title and filename:
                title = Path(filename).stem
            
            # Extract digest from first paragraph
            digest = ""
            for line in lines[1:]:  # Skip title line
                line = line.strip()
                if line and len(line) > 20:
                    digest = line
                    break
            
            # Convert to HTML (simple paragraph conversion)
            paragraphs = full_text.split('\n\n')
            html_parts = []
            for para in paragraphs:
                para = para.strip()
                if para:
                    # Escape HTML entities
                    para = para.replace('&', '&amp;')
                    para = para.replace('<', '&lt;')
                    para = para.replace('>', '&gt;')
                    # Convert line breaks within paragraph
                    para = para.replace('\n', '<br>')
                    html_parts.append(f'<p>{para}</p>')
            
            html_content = '\n'.join(html_parts)
            
            return title[:64], digest[:200], html_content
            
    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"PDF 解析失败: {str(e)}")


def parse_docx(content: bytes, filename: str = "") -> Tuple[str, str, str]:
    """
    Parse Word document (.docx) to HTML.
    
    Args:
        content: Raw file content
        filename: Original filename
        
    Returns:
        Tuple of (title, digest, html_content)
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise DocumentParseError("缺少 python-docx 库，请安装: pip install python-docx")
    
    try:
        # Open document from bytes
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        
        if not doc.paragraphs:
            raise DocumentParseError("Word 文档为空")
        
        # Extract title from first heading or first paragraph
        title = ""
        digest = ""
        html_parts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading
            style_name = para.style.name.lower() if para.style else ""
            
            if 'heading' in style_name or 'title' in style_name:
                level = 1
                if 'heading 2' in style_name:
                    level = 2
                elif 'heading 3' in style_name:
                    level = 3
                
                if not title and level == 1:
                    title = text
                
                html_parts.append(f'<h{level}>{_escape_html(text)}</h{level}>')
            else:
                # Regular paragraph
                if not digest and len(text) > 20:
                    digest = text
                
                # Check for bold/italic runs
                html_text = _convert_runs_to_html(para)
                html_parts.append(f'<p>{html_text}</p>')
        
        # Use filename as title if not found
        if not title and filename:
            title = Path(filename).stem
        
        # Handle tables
        for table in doc.tables:
            html_parts.append(_convert_table_to_html(table))
        
        html_content = '\n'.join(html_parts)
        
        return title[:64], digest[:200], html_content
        
    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"Word 文档解析失败: {str(e)}")


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;'))


def _convert_runs_to_html(para) -> str:
    """Convert paragraph runs to HTML with formatting."""
    html_parts = []
    
    for run in para.runs:
        text = run.text
        if not text:
            continue
        
        text = _escape_html(text)
        
        # Apply formatting
        if run.bold:
            text = f'<strong>{text}</strong>'
        if run.italic:
            text = f'<em>{text}</em>'
        if run.underline:
            text = f'<u>{text}</u>'
        
        html_parts.append(text)
    
    return ''.join(html_parts) if html_parts else _escape_html(para.text)


def _convert_table_to_html(table) -> str:
    """Convert Word table to HTML."""
    html = ['<table border="1" style="border-collapse: collapse;">']
    
    for row in table.rows:
        html.append('<tr>')
        for cell in row.cells:
            text = _escape_html(cell.text.strip())
            html.append(f'<td style="padding: 8px;">{text}</td>')
        html.append('</tr>')
    
    html.append('</table>')
    return '\n'.join(html)


# Supported file extensions and their parsers
SUPPORTED_FORMATS = {
    '.md': ('Markdown', parse_markdown),
    '.markdown': ('Markdown', parse_markdown),
    '.pdf': ('PDF', parse_pdf),
    '.docx': ('Word', parse_docx),
}


def get_supported_formats() -> list:
    """Get list of supported document formats."""
    return [
        {'extension': ext, 'name': name}
        for ext, (name, _) in SUPPORTED_FORMATS.items()
    ]


def parse_document(content: bytes, filename: str) -> Tuple[str, str, str]:
    """
    Parse document based on file extension.
    
    Args:
        content: Raw file content
        filename: Original filename with extension
        
    Returns:
        Tuple of (title, digest, html_content)
        
    Raises:
        DocumentParseError: If format is not supported or parsing fails
    """
    ext = Path(filename).suffix.lower()
    
    if ext not in SUPPORTED_FORMATS:
        supported = ', '.join(SUPPORTED_FORMATS.keys())
        raise DocumentParseError(f"不支持的文件格式: {ext}。支持的格式: {supported}")
    
    _, parser = SUPPORTED_FORMATS[ext]
    return parser(content, filename)
